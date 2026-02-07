import os
import json
import sys
import logging

logger = logging.getLogger(__name__)

# region 辅助类


class MarkdownAstParser:
    """
    Markdown AST 解析器类
    封装 markdown-it 库的功能，用于将 Markdown 文本解析为 token 流
    """

    def __init__(self):
        """
        初始化 Markdown 解析器实例
        """
        # 延迟初始化 MarkdownIt，只有在 parse 时才导入并创建实例
        self.md = None

    def parse(self, md_text: str):
        """
        解析 Markdown 文本为 token 列表

        Args:
            md_text (str): 输入的 Markdown 文本内容

        Returns:
            list: 解析后的 token 列表
        """
        logger.info("正在解析Markdown文档结构...")
        if self.md is None:
            from markdown_it import MarkdownIt

            self.md = MarkdownIt()
        return self.md.parse(md_text)


class ListNode:
    """
    列表节点类
    用于构建和存储 Markdown 列表的树状结构
    """

    def __init__(self, content="", ordered=False, level=0):
        """
        初始化列表节点

        Args:
            content (str): 列表项的内容文本
            ordered (bool): 是否为有序列表
            level (int): 列表嵌套层级（从0开始）
        """
        self.content = content
        self.ordered = ordered
        self.level = level
        self.children = []


# endregion


class MdToDocx:
    """
    Markdown 转 Word 文档转换器核心类
    负责处理 Markdown 解析、样式应用和 Word 文档生成
    """

    def __init__(self):
        """
        初始化转换器
        创建 AST 解析器实例
        """
        self.parser = MarkdownAstParser()

    def _ensure_docx(self):
        """
        延迟导入 python-docx 相关模块并绑定到实例属性，避免模块导入期间的开销。
        """
        if getattr(self, "_docx_loaded", False):
            return

        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.oxml.ns import qn
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
        from docx.oxml import OxmlElement

        self.Document = Document
        self.Inches = Inches
        self.Pt = Pt
        self.RGBColor = RGBColor
        self.qn = qn
        self.WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH
        self.WD_LINE_SPACING = WD_LINE_SPACING
        self.OxmlElement = OxmlElement
        self._docx_loaded = True

    def convert(self, md_text: str, output_path: str, styles: list = None):
        """
        执行转换流程：Markdown -> Word

        Args:
            md_text (str): 原始 Markdown 文本
            output_path (str): 输出 Word 文档的路径 (.docx)
            styles (list, optional): 自定义样式配置列表. Defaults to None.
        """
        # 确保按需加载 python-docx
        self._ensure_docx()

        if os.path.exists(output_path):
            self.doc = self.Document(output_path)
        else:
            self.doc = self.Document()

        self._enable_doc_grid()

        tokens = self.parser.parse(md_text)
        # 获取默认样式
        if styles is None:
            styles = []
            current_dir = os.path.dirname(os.path.abspath(sys.argv[0]))
            with open(
                os.path.join(current_dir, "style.json"), "r", encoding="utf-8"
            ) as f:
                styles = json.load(f)
                styles = styles.get("default", [])
        self.styles = styles

        logger.info("正在写入Word文档...")
        self._write_tokens(tokens)
        logger.success("Word文档写入完毕！")
        self.doc.save(output_path)

    # region 写入内容

    def _write_tokens(self, tokens, paragraph=None, paragraph_style=None):
        """
        遍历并处理 Markdown token 流，将其转换为 Word 文档元素

        Args:
            tokens (list): token 列表
            paragraph (docx.text.paragraph.Paragraph, optional): 当前正在处理的段落对象. Defaults to None.
            paragraph_style (str, optional): 当前段落应用的样式名称. Defaults to None.
        """
        i = 0
        while i < len(tokens):
            t = tokens[i]

            # 解析标题级别
            if t.type == "heading_open":
                self._current_heading_level = int(t.tag[1])

            # 解析标题内容
            elif t.type == "inline" and hasattr(self, "_current_heading_level"):
                if paragraph is None:
                    paragraph = self.doc.add_heading(
                        t.content, level=self._current_heading_level
                    )
                else:
                    paragraph.add_run(t.content)
                # 应用标题样式
                self._set_paragraph_style(paragraph, f"h{self._current_heading_level}")
                paragraph = None
                # 重置标题级别
                del self._current_heading_level

            # 列表处理
            elif t.type in ("bullet_list_open", "ordered_list_open"):
                # 先解析当前列表为树结构
                nodes, next_i = self._parse_list(tokens, i)
                # 再写入 Word
                self._write_list_to_word(nodes)
                i = next_i
                continue  # 跳过 i += 1

            # 普通段落 + 图片
            elif t.type == "inline":
                self._handle_inline(t, paragraph, paragraph_style)

            i += 1

    def _handle_inline(self, token, paragraph=None, paragraph_style=None):
        """
        处理行内元素（如普通文本、图片等）

        Args:
            token (Token): 当前处理的 inline token
            paragraph (docx.text.paragraph.Paragraph, optional): 目标段落对象
            paragraph_style (str, optional): 样式名称
        """
        if paragraph_style is None:
            paragraph_style = "text"

        for child in token.children:
            if paragraph is None:
                paragraph = self.doc.add_paragraph()

            # ✅ 图片
            if child.type == "image":
                src = child.attrs.get("src")
                self._add_image(src, paragraph)

            # ✅ 普通文本
            elif child.type == "text" and child.content.strip() != "":
                paragraph.add_run(child.content)
                self._set_paragraph_style(paragraph, paragraph_style)

    def _add_image(self, src: str, paragraph=None):
        """
        向段落添加图片，并自动计算合适的显示宽度
        支持列表内图片宽度正确计算

        Args:
            src (str): 图片来源 (路径/URL/Base64)
            paragraph (docx.text.paragraph.Paragraph, optional): 目标段落. Defaults to None.
        """

        # 延迟导入 requests/base64/tempfile 等
        self._ensure_docx()
        from urllib.parse import unquote
        import requests
        import base64
        import tempfile

        # ---------- 1️⃣ 获取图片 ----------
        if src.startswith("http"):
            src = unquote(src)
            r = requests.get(src, timeout=10)
            r.raise_for_status()
            img_bytes = r.content

        elif src.startswith("data:image"):
            img_bytes = base64.b64decode(src.split(",", 1)[1])

        else:
            src = unquote(src)
            with open(src, "rb") as f:
                img_bytes = f.read()

        with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmp:
            tmp.write(img_bytes)
            tmp_path = tmp.name

        # ---------- 2️⃣ paragraph ----------
        if paragraph is None:
            paragraph = self.doc.add_paragraph()

        # ---------- 3️⃣ 计算宽度 ----------
        section = self.doc.sections[0]

        page_width = section.page_width
        margin_left = section.left_margin
        margin_right = section.right_margin

        # 🔥 核心：列表层级缩进
        list_level = self._get_list_level(paragraph)
        list_indent = self.Inches(0.25 * (list_level + 1))

        available_width = page_width - margin_left - margin_right - list_indent

        # 防御
        if isinstance(available_width, int):
            available_width = self.Inches(available_width / 914400)

        try:
            inches_val = available_width.inches
        except Exception:
            inches_val = None

        if inches_val is None or (
            isinstance(inches_val, (int, float)) and inches_val < 1
        ):
            available_width = self.Inches(1)

        # ---------- 4️⃣ 插入 ----------
        # 如果段落已有多个 run，先追加换行
        if len(paragraph.runs) > 0:
            paragraph.add_run("\n")
        run = paragraph.add_run()
        run.add_picture(tmp_path, width=available_width)

    # endregion

    # region 列表处理

    def _parse_list(self, tokens, start=0, level=0):
        """
        解析列表 tokens 为 ListNode 树状结构 (递归)

        Args:
            tokens (list): token 列表
            start (int): 当前解析起始索引
            level (int): 当前列表层级

        Returns:
            tuple: (解析得到的节点列表, 下一个处理索引)
        """
        nodes = []
        i = start
        ordered = tokens[start].type == "ordered_list_open"

        while i < len(tokens):
            t = tokens[i]

            if t.type == "list_item_open":
                node = ListNode(level=level, ordered=ordered)
                j = i + 1
                while j < len(tokens) and tokens[j].type != "list_item_close":
                    if tokens[j].type == "inline":
                        node.content = tokens[j].content
                    # 子列表
                    elif tokens[j].type in ("bullet_list_open", "ordered_list_open"):
                        child_nodes, next_j = self._parse_list(tokens, j, level + 1)
                        node.children.extend(child_nodes)
                        j = next_j - 1  # 解析完子列表后更新 j
                    j += 1
                nodes.append(node)
                i = j

            elif t.type in ("bullet_list_open", "ordered_list_open"):
                # 外层列表开始，更新 ordered
                ordered = t.type == "ordered_list_open"

            elif t.type in ("bullet_list_close", "ordered_list_close"):
                return nodes, i + 1

            i += 1

        return nodes, i

    def _write_list_to_word(self, nodes):
        """
        将列表节点树写入 Word 文档

        Args:
            nodes (list[ListNode]): 列表节点集合
        """

        def _write_nodes(nodes, level=0):
            for node in nodes:
                style = self._get_list_style(node.ordered, level)

                paragraph = self.doc.add_paragraph(style=style)
                self._set_paragraph_style(paragraph, f"li")
                tokens = self.parser.parse(node.content)
                self._write_tokens(tokens, paragraph=paragraph, paragraph_style=f"li")

                # 递归写入子列表
                if node.children:
                    _write_nodes(node.children, level + 1)

        _write_nodes(nodes)

    def _get_list_style(self, ordered, level):
        """
        根据列表类型和层级获取 Word 列表样式名

        Args:
            ordered (bool): 是否有序
            level (int): 层级

        Returns:
            str: 样式名称
        """
        if ordered:
            styles = [
                "List Number",
                "List Number 2",
                "List Number 3",
            ]
        else:
            styles = [
                "List Bullet",
                "List Bullet 2",
                "List Bullet 3",
            ]

        # Word 默认只内置到 3 级，超过就复用最后一级
        return styles[min(level, len(styles) - 1)]

    def _get_list_level(self, paragraph):
        """
        获取段落的列表层级

        Args:
            paragraph (docx.text.paragraph.Paragraph): 段落对象

        Returns:
            int: 列表层级 (从0开始), 非列表返回 -1
        """
        p = paragraph._p
        pPr = p.pPr

        # 1️⃣ OOXML 编号列表
        if pPr is not None and pPr.numPr is not None:
            ilvl = pPr.numPr.ilvl
            if ilvl is not None:
                return int(ilvl.val)

        # 2️⃣ 样式列表（List Bullet / List Number）
        style = paragraph.style
        if style and style.name:
            name = style.name.lower()
            if name.startswith("list"):
                # 从样式名中提取数字作为层级
                import re

                match = re.search(r"(\d+)", name)
                return int(match.group(1)) - 1 if match else 0

        # 3️⃣ 非列表
        return -1

    # endregion

    # region 样式设置

    def _set_paragraph_style(self, paragraph, style_name):
        """
        应用段落和字体样式
        完整版：
        - 支持任意倍数行距下的“视觉垂直居中”
        - 与人工 Word/WPS 文档行为一致
        - 不破坏列表（li）编号结构

        Args:
            paragraph (docx.text.paragraph.Paragraph): 目标段落
            style_name (str): 样式配置名称
        """
        # 确保 docx 系列符号已按需加载
        self._ensure_docx()
        style = self.styles.get(style_name)
        if not style:
            return

        # =====================================================
        # 1️⃣ 段落级格式（pPr）
        # =====================================================
        p_format = paragraph.paragraph_format
        pPr = paragraph._element.get_or_add_pPr()

        # ---- textAlignment = auto（与人工文档一致）
        text_align = pPr.find(self.qn("w:textAlignment"))
        if text_align is None:
            text_align = self.OxmlElement("w:textAlignment")
            pPr.append(text_align)
        text_align.set(self.qn("w:val"), "auto")

        # ---- snapToGrid = 1（关键：启用基线网格）
        snap = pPr.find(self.qn("w:snapToGrid"))
        if snap is None:
            snap = self.OxmlElement("w:snapToGrid")
            pPr.append(snap)
        snap.set(self.qn("w:val"), "1")

        # ---- 中文排版辅助属性（不影响西文）
        for tag in ["w:kinsoku", "w:overflowPunct", "w:adjustRightInd"]:
            if pPr.find(self.qn(tag)) is None:
                pPr.append(self.OxmlElement(tag))

        # ---- 对齐方式
        if "align" in style:
            align_map = {
                "left": self.WD_ALIGN_PARAGRAPH.LEFT,
                "center": self.WD_ALIGN_PARAGRAPH.CENTER,
                "right": self.WD_ALIGN_PARAGRAPH.RIGHT,
                "justify": self.WD_ALIGN_PARAGRAPH.JUSTIFY,
            }
            p_format.alignment = align_map.get(
                style["align"], self.WD_ALIGN_PARAGRAPH.LEFT
            )

        # ---- 行距（不固定倍数）
        if "line_spacing" in style:
            p_format.line_spacing = style["line_spacing"]
            p_format.line_spacing_rule = self.WD_LINE_SPACING.MULTIPLE

        p_format.space_before = self.Pt(style.get("space_before", 0))
        p_format.space_after = self.Pt(style.get("space_after", 0))

        # =====================================================
        # 2️⃣ 缩进规则
        # =====================================================
        if "first_line_indent" in style:
            p_format.first_line_indent = self.Pt(
                style.get("font_size", 11) * style["first_line_indent"]
            )

        # 🔥 li 特殊处理
        if style_name == "li":
            p_format.left_indent = None
            p_format.first_line_indent = None

        # =====================================================
        # 3️⃣ 字符级（rPr）
        # =====================================================
        font_name = style.get("font_name", "微软雅黑")
        font_size = style.get("font_size", 11)

        for run in paragraph.runs:
            font = run.font
            rPr = run._element.get_or_add_rPr()

            # ---- 字体四槽位（ascii / hAnsi / eastAsia / cs），
            # 解决“中文字体不生效”问题
            rFonts = rPr.get_or_add_rFonts()
            rFonts.set(self.qn("w:ascii"), font_name)
            rFonts.set(self.qn("w:hAnsi"), font_name)
            rFonts.set(self.qn("w:eastAsia"), font_name)
            rFonts.set(self.qn("w:cs"), font_name)

            # ---- 语言环境
            lang = rPr.find(self.qn("w:lang"))
            if lang is None:
                lang = self.OxmlElement("w:lang")
                rPr.append(lang)
            lang.set(self.qn("w:val"), "en-US")
            lang.set(self.qn("w:eastAsia"), "zh-CN")

            # ---- 基础样式
            font.size = self.Pt(font_size)
            font.bold = style.get("bold", False)
            font.italic = style.get("italic", False)
            font.underline = style.get("underline", False)

            if "font_color" in style:
                c = style["font_color"].lstrip("#")
                font.color.rgb = self.RGBColor(
                    int(c[0:2], 16),
                    int(c[2:4], 16),
                    int(c[4:6], 16),
                )

    def _enable_doc_grid(self):
        """
        启用 Word 中文排版网格，实现段落文字垂直居中，只需要执行一次即可
        - 始终确保 w:type="lines"
        - 不强制 linePitch，由 Word 自动计算
        """
        # 确保 docx 相关对象已加载
        self._ensure_docx()
        section = self.doc.sections[0]
        sectPr = section._sectPr

        # 查找已有 docGrid
        docGrids = sectPr.xpath(".//w:docGrid")
        if docGrids:
            docGrid = docGrids[0]
        else:
            docGrid = self.OxmlElement("w:docGrid")
            sectPr.append(docGrid)

        # ✅ 核心：始终设置为 lines
        docGrid.set(self.qn("w:type"), "lines")

        # 可选：显式关闭字符网格（推荐）
        docGrid.set(self.qn("w:charSpace"), "0")

    # endregion
